from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from PyPDF2 import PdfReader


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt"}


def _load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def _load_docx(path: Path) -> str:
    paragraphs = [paragraph.text.strip() for paragraph in DocxDocument(path).paragraphs]
    return "\n".join(text for text in paragraphs if text).strip()


def _load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    texts = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n".join(text for text in texts if text).strip()


def load_document(path: Path) -> Document:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported file type: {suffix or '<none>'}")

    loaders = {
        ".txt": _load_txt,
        ".docx": _load_docx,
        ".pdf": _load_pdf,
    }
    content = loaders[suffix](path)
    return Document(
        page_content=content,
        metadata={"source": str(path), "file_type": suffix.removeprefix(".")},
    )


def load_documents(paths: list[Path]) -> list[Document]:
    documents: list[Document] = []
    for path in paths:
        document = load_document(path)
        if document.page_content.strip():
            documents.append(document)
    return documents
